import os
import time
import telebot
import google.generativeai as genai
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- بيانات الوصول المباشرة ---
TELEGRAM_TOKEN = "8433118363:AAH0iqeZVo3xz-KP_KQ7LxHSdhRZnOmb2LQ"
GEMINI_API_KEY = "AIzaSyCChd6IL-8hi9ttKOIwH-vVF57MzK8X26s"

# إعداد محرك الذكاء الاصطناعي (نسخة المطور قيس)
genai.configure(api_key=GEMINI_API_KEY)
ai_config = {
    "temperature": 0.75,
    "top_p": 0.9,
    "max_output_tokens": 4096,
}

# تعليمات النظام لتعزيز شخصية البوت
SYSTEM_PROMPT = (
    "أنت نظام ذكاء اصطناعي متطور جداً، تم إنشاؤك وتطويرك بواسطة المطور (قيس). "
    "مهمتك الأساسية هي ترجمة النصوص من الإنجليزية إلى العربية بدقة متناهية وأسلوب بلاغي رائع. "
    "يجب أن تظهر الترجمة كأنها مكتوبة بيد خبير، مع الحفاظ على روح النص الأصلي."
)

ai_model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=ai_config,
    system_instruction=SYSTEM_PROMPT
)

bot = telebot.TeleBot(TELEGRAM_TOKEN)

# دالة الترجمة المحمية من القلتشات
def perform_smart_translation(text_content):
    if not text_content or len(text_content.strip()) < 2:
        return None
    
    attempts = 0
    while attempts < 3:
        try:
            prompt = f"Translate the following English text into professional, fluent Arabic:\n\n{text_content}"
            response = ai_model.generate_content(prompt)
            if response and response.text:
                return response.text.strip()
            return None
        except Exception as error:
            attempts += 1
            time.sleep(2)
    return None

# رسالة الترحيب - بالعربي أولاً كما طلبت
@bot.message_handler(commands=['start', 'help'])
def send_welcome_message(message):
    welcome_text = (
        "مرحباً بك! أنا نظام ترجمة الملفات الذكي، نسخة مطورة ومحدثة بالكامل من قبل المطور **قيس**.\n\n"
        "Welcome! I am a smart file translation system, a fully developed and updated version by the developer **Qais**.\n\n"
        "--- المميزات / Features ---\n"
        "✅ الحفاظ على الصور والجداول / Preserves images and tables.\n"
        "✅ ترجمة تحت كل سطر / Interlinear translation.\n"
        "✅ تنسيق لوني مريح / Comfortable color formatting.\n\n"
        "أرسل ملف .docx للبدء / Send a .docx file to start."
    )
    bot.reply_to(message, welcome_text, parse_mode="Markdown")

# معالجة المستندات
@bot.message_handler(content_types=['document'])
def handle_incoming_document(message):
    file_name = message.document.file_name
    
    # فحص الامتداد
    if not file_name.lower().endswith('.docx'):
        bot.reply_to(message, "❌ Please send a .docx file only.")
        return

    # رسالة الانتظار بالإنجليزي كما طلبت
    progress_msg = bot.reply_to(message, "Please wait...")
    
    # إنشاء أسماء ملفات فريدة لتجنب التداخل
    timestamp = int(time.time())
    input_path = f"in_{timestamp}_{file_name}"
    output_path = f"translated_by_Qais_{file_name}"

    try:
        # تحميل الملف من سيرفرات تلغرام
        file_raw = bot.get_file(message.document.file_id)
        file_bytes = bot.download_file(file_raw.file_path)
        
        with open(input_path, 'wb') as f:
            f.write(file_bytes)

        # فتح المستند للبدء في التعديل
        document = Document(input_path)
        
        # 1. معالجة الفقرات العادية
        for paragraph in document.paragraphs:
            original_text = paragraph.text.strip()
            if len(original_text) > 3:
                translated_text = perform_smart_translation(original_text)
                if translated_text:
                    # إضافة سطر جديد للنص المترجم
                    run = paragraph.add_run(f"\n{translated_text}")
                    # تنسيق لون الترجمة (أزرق سماوي غامق مريح للنظر)
                    run.font.color.rgb = RGBColor(31, 73, 125)
                    run.font.bold = True
                    run.font.size = Pt(11)

        # 2. معالجة النصوص داخل الجداول (لأنها غالباً ما تحتوي بيانات مهمة)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if len(paragraph.text.strip()) > 2:
                            translated_text = perform_smart_translation(paragraph.text)
                            if translated_text:
                                run = paragraph.add_run(f"\n{translated_text}")
                                run.font.color.rgb = RGBColor(31, 73, 125)
                                run.font.italic = True

        # حفظ التعديلات
        document.save(output_path)

        # إرسال الملف النهائي للمستخدم
        with open(output_path, 'rb') as final_file:
            # رسالة التم بالإنجليزي كما طلبت
            bot.send_document(
                message.chat.id, 
                final_file, 
                caption="Done! Translated by Qais AI System."
            )

    except Exception as global_error:
        bot.reply_to(message, f"An error occurred: {str(global_error)}")
    
    finally:
        # نظام التنظيف الذكي للملفات لتجنب امتلاء الذاكرة في Render
        for temp_file in [input_path, output_path]:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass
        # حذف رسالة "Please wait" بعد الانتهاء
        try:
            bot.delete_message(message.chat.id, progress_msg.message_id)
        except:
            pass

# تشغيل البوت مع نظام الحماية من التوقف المفاجئ
print("Qais AI Translation Bot is now Active...")
while True:
    try:
        bot.polling(none_stop=True, interval=0, timeout=40)
    except Exception as e:
        print(f"Reconnect Log: {e}")
        time.sleep(10)
